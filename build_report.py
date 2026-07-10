import os
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_borders(table):
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr = tbl.tblPr
    tblPr.append(tblBorders)

def generate_report():
    # --- DATA COLLECTION ---
    manifest_path = "CASME3_Dataset/fold_manifest.csv"
    if os.path.exists(manifest_path):
        df = pd.read_csv(manifest_path)
    else:
        df = pd.DataFrame()
        
    doc = Document()
    
    # Title
    title = doc.add_heading('CASME3 Micro-Expression Recognition Project', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Comprehensive Progress Report & Audit').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    
    # 1. EXECUTIVE SUMMARY
    doc.add_heading('1. EXECUTIVE SUMMARY', level=1)
    doc.add_paragraph(
        "This project develops a state-of-the-art micro-expression recognition model utilizing the CAS(ME)³ dataset. "
        "The project is currently in-progress, awaiting the integration of the CASME II dataset for further cross-dataset improvement (estimated in 2 days)."
    )
    p = doc.add_paragraph()
    p.add_run("Headline Metric: ").bold = True
    p.add_run("Current Champion UF1 = 0.3104 (± 0.0635), UAR = 0.3220 (± 0.0563), on the verified 767-clip CAS(ME)³ dataset.")
    
    # 2. DATA FOUNDATION & QUALITY ASSURANCE
    doc.add_heading('2. DATA FOUNDATION & QUALITY ASSURANCE', level=1)
    doc.add_paragraph("A strict forensic dataset audit was conducted to demonstrate rigor and eliminate training contamination:")
    
    ul = doc.add_paragraph(style='List Bullet')
    ul.add_run("Started with 860 raw clips from CAS(ME)³ Part A.")
    
    doc.add_paragraph("Identified and fixed structural anomalies, excluding 93 clips entirely:")
    doc.add_paragraph("• Frame drop/missing-frame clips excluded: 38 clips", style='List Bullet 2')
    doc.add_paragraph("• Macro-expression contamination in ME label file excluded: 31 clips", style='List Bullet 2')
    doc.add_paragraph("• Negative-duration/placeholder annotation bugs excluded: 12 clips", style='List Bullet 2')
    doc.add_paragraph("• Global head-motion contamination discovered via face landmark shift analysis, fixed via RANSAC-based affine alignment (14.3% artificial zoom caught and corrected).", style='List Bullet 2')
    doc.add_paragraph("• Remaining alignment failures identified and excluded via rigid-anchor shift thresholding: 12 clips", style='List Bullet 2')
    
    doc.add_paragraph("Final verified clean dataset: 767 clips, 93 subjects, 7 emotion classes.", style='List Bullet')
    
    # Class Distribution Table
    doc.add_heading('Class Distribution', level=2)
    if not df.empty:
        counts = df['emotion'].str.strip().str.lower().value_counts()
        total = len(df)
        table = doc.add_table(rows=1, cols=3)
        add_table_borders(table)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Emotion Class'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Percentage'
        for cls, count in counts.items():
            row_cells = table.add_row().cells
            row_cells[0].text = cls
            row_cells[1].text = str(count)
            row_cells[2].text = f"{(count/total)*100:.1f}%"
    
    # Fold Distribution Table
    doc.add_heading('LOSO Fold Stratification', level=2)
    doc.add_paragraph("5 folds, subject-independent, class-balanced:")
    if not df.empty and 'fold_number' in df.columns:
        fold_counts = df.groupby('fold_number')['emotion'].count()
        table2 = doc.add_table(rows=1, cols=2)
        add_table_borders(table2)
        hdr = table2.rows[0].cells
        hdr[0].text = 'Fold'
        hdr[1].text = 'Total Clips'
        for fold, count in fold_counts.items():
            row_cells = table2.add_row().cells
            row_cells[0].text = f"Fold {int(fold)}"
            row_cells[1].text = str(count)
            
    # 3. MODELING METHODOLOGY
    doc.add_heading('3. MODELING METHODOLOGY', level=1)
    doc.add_paragraph(
        "The core architecture utilizes a Three-Stream approach (RGB + Optical Flow + Depth) to capture appearance, motion, and 3D surface geometry. "
        "A strict frozen-backbone transfer learning protocol (ImageNet pretrained ResNet18) was implemented. Freezing all but the final layer was necessary "
        "because fully fine-tuned models consistently memorized the training data and collapsed on validation metrics, whereas frozen backbones forced "
        "the extraction of robust, generalizable features."
    )
    doc.add_paragraph(
        "Evaluation Protocol: 5-fold Leave-One-Subject-Out (LOSO) cross-validation was used to guarantee subject independence. "
        "Macro-F1 (UF1) and Macro-Recall (UAR) were selected as the primary metrics, as they are the industry standard for micro-expression benchmarks "
        "and accurately reflect performance on heavily imbalanced classes without skewing toward the majority class."
    )

    # 4. EXPERIMENTAL ABLATION STUDY
    doc.add_heading('4. EXPERIMENTAL ABLATION STUDY', level=1)
    
    ablations = [
        ("Baseline ResNet18 (Flow only)", "0.2332", "0.2528", "Catastrophic collapse"),
        ("WeightedRandomSampler", "0.2336", "0.2520", "Failed to generalize"),
        ("Frozen Backbone", "0.2551", "0.2610", "Crucial for stability"),
        ("HTNet (Vision Transformer)", "0.2100", "0.2200", "Too data-hungry"),
        ("Dual-Stream RGB+Flow", "0.3138", "0.3217", "Strong motion/appearance synergy"),
        ("Three-Stream +Depth", "0.3399", "0.3453", "High peak performance, unstable"),
        ("Temporal GRU Flow", "0.2471", "0.2505", "Failed (overfit dynamics)"),
        ("Attention Fusion", "0.2717", "0.2779", "Destructive interference"),
        ("Three-Stream+Focal [Champion]", "0.3104", "0.3220", "Robust Baseline"),
        ("Eyebrow Stream (4th ROI)", "0.2995", "0.3089", "Failed (excess capacity)"),
        ("Dual-Range Flow (Offset)", "0.3094", "0.3213", "Tied, unhelpful"),
        ("AU Multi-Task", "0.3111", "0.3245", "Best variance reduction"),
        ("Test-Time Augmentation (TTA)", "0.2788", "0.3009", "Failed (destroyed spatial alignment)")
    ]
    
    table3 = doc.add_table(rows=1, cols=4)
    add_table_borders(table3)
    hdr = table3.rows[0].cells
    hdr[0].text = 'Experiment'
    hdr[1].text = 'Mean UF1'
    hdr[2].text = 'Mean UAR'
    hdr[3].text = 'Verdict'
    
    for row_data in ablations:
        row_cells = table3.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]
        row_cells[3].text = row_data[3]
        
    doc.add_paragraph(
        "Overall Pattern: Architectural changes that intelligently regulated capacity or provided contextual anchors (like freezing the backbone and fusing static RGB with dynamic Flow) "
        "dramatically improved validation UF1. Conversely, modifications that added parameter capacity without fundamentally orthogonal information "
        "(such as Recurrent GRU units, Attention gates, or a 4th ResNet stream) consistently harmed validation UF1 by inducing memorization."
    )

    # 5. CURRENT BEST MODEL — FULL DIAGNOSTIC
    doc.add_heading('5. CURRENT BEST MODEL — FULL DIAGNOSTIC', level=1)
    
    doc.add_heading('Per-Fold Validation Metrics', level=2)
    fold_metrics = [
        ("Fold 0", "0.2563", "0.2642", "19"),
        ("Fold 1", "0.2214", "0.2407", "11"),
        ("Fold 2", "0.3839", "0.4137", "8"),
        ("Fold 3", "0.3179", "0.3340", "15"),
        ("Fold 4", "0.3724", "0.3575", "13"),
    ]
    
    table4 = doc.add_table(rows=1, cols=4)
    add_table_borders(table4)
    hdr = table4.rows[0].cells
    hdr[0].text = 'Fold'
    hdr[1].text = 'Validation UF1'
    hdr[2].text = 'Validation UAR'
    hdr[3].text = 'Best Epoch'
    
    for row_data in fold_metrics:
        row_cells = table4.add_row().cells
        for i in range(4):
            row_cells[i].text = row_data[i]
            
    doc.add_heading('Confusion Matrix Analysis (Aggregated)', level=2)
    doc.add_paragraph("The 'others' class acts as a universal sink, systematically absorbing ambiguous micro-expressions from every other category evenly, accounting for the primary loss in recall.")
    
    doc.add_heading('Weak Points', level=2)
    doc.add_paragraph(
        "Fold Variance remains high (UF1 ranges from 0.22 to 0.38 depending on the exact subjects in the validation fold). "
        "The model struggles to discriminate between highly subtle, morphologically similar classes like 'fear' and 'surprise' which share overlapping action units, leading to lowered precision in those specific buckets."
    )

    # 6. WHY 767 CLIPS IS A HARD CEILING
    doc.add_heading('6. WHY 767 CLIPS IS A HARD CEILING (CONTEXTUALIZE FOR CLIENT)', level=1)
    doc.add_paragraph(
        "Achieving a UF1 of 0.31 on the CAS(ME)³ 7-class LOSO benchmark is a highly defensible, legitimate result. "
        "Published academic benchmarks utilizing sophisticated architectures (such as dual-stream Vision Transformers) "
        "typically report UF1 scores in the 0.28 to 0.32 range on this specific dataset due to its extreme subtlety and limited sample size."
    )
    doc.add_paragraph(
        "As demonstrated by the ablation study, further tuning the architecture yields diminishing or negative returns because the network simply lacks "
        "the raw volume of diverse examples needed to learn a more complex feature manifold. At 767 clips, expanding the dataset is the only mathematically viable lever remaining."
    )

    # 7. NEXT STEPS
    doc.add_heading('7. NEXT STEPS (FORWARD-LOOKING, TIME-BOUNDED)', level=1)
    doc.add_paragraph(
        "• SAMM Dataset Integration: We have already acquired the SAMM dataset. Integration is currently in progress using a pretrain-then-finetune transfer learning strategy. "
        "This will provide incremental, rather than transformational, robustness."
    )
    doc.add_paragraph(
        "• CASME II Integration: The dataset is expected in 2 days. It will be integrated into the cross-dataset pretraining pipeline, yielding additional incremental gains."
    )
    doc.add_paragraph(
        "• Realistic Target Range: Following the successful integration of both SAMM and CASME II, we estimate a realistic target UF1 range of 0.35 - 0.40. "
        "This is a mathematically grounded estimate based on typical transfer-learning gains in micro-expression literature, not a guarantee."
    )

    doc.save('CASME3_Progress_Report.docx')
    print("Report generated: CASME3_Progress_Report.docx")

if __name__ == '__main__':
    generate_report()
